from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/shubhamsingh/Documents/New project/homepage/outputs/WCEM-Proposal-Letter.docx"

NAVY = "152A43"
INK = "1C2B39"
SLATE = "5E6C7B"
MUTED = "7E8A97"
GOLD = "9A7A37"
LINE = "D7E0E8"
SOFT = "F7F9FC"
WHITE = "FFFFFF"


def set_font(run, name, size, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        child = tc_mar.find(qn(f"w:{key}"))
        if child is None:
            child = OxmlElement(f"w:{key}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_table_borders(table, color=WHITE, size="0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, edge_data in kwargs.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def add_rule(paragraph, color=LINE, size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, "Times New Roman", 9.0, color=INK)
    set_paragraph_format(p, after=2, line=1.0)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.38)
section.bottom_margin = Inches(0.36)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.header_distance = Inches(0.18)
section.footer_distance = Inches(0.18)
section.different_first_page_header_footer = False
section.start_type = WD_SECTION.NEW_PAGE

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(11)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp.add_run("SrS LOGICS")
set_font(r, "Times New Roman", 11, bold=True, color=NAVY)
r = hp.add_run("  |  srslogics.com")
set_font(r, "Times New Roman", 9.4, color=SLATE)
set_paragraph_format(hp, after=1)

hp2 = header.add_paragraph()
hp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp2.add_run("Founder-led software engineering company | Nagpur, India | +91 9270925106 | shubhamsingh@srslogics.com")
set_font(r, "Times New Roman", 8.6, color=SLATE)
set_paragraph_format(hp2, after=1)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = date_p.add_run("20 June 2026")
set_font(r, "Times New Roman", 8.8, bold=False, color=SLATE)
set_paragraph_format(date_p, after=2)

p = doc.add_paragraph()
r = p.add_run("Strategic Proposal for Unified Digital Institution Transformation")
set_font(r, "Times New Roman", 13.2, bold=True, color=NAVY)
set_paragraph_format(p, after=2, line=1.0)

p = doc.add_paragraph()
r = p.add_run(
    "For Wainganga College of Engineering and Management (WCEM), Nagpur."
)
set_font(r, "Times New Roman", 8.8, color=SLATE)
set_paragraph_format(p, after=2, line=1.0)

p = doc.add_paragraph()
set_paragraph_format(p, after=3)
add_rule(p)

p = doc.add_paragraph()
r = p.add_run("To,\nThe Director,\nWainganga College of Engineering and Management (WCEM),\nNagpur.")
set_font(r, "Times New Roman", 8.8, color=INK)
set_paragraph_format(p, after=3, line=1.0)

p = doc.add_paragraph()
r = p.add_run("Subject: Strategic Proposal for Unified Digital Institution Transformation")
set_font(r, "Times New Roman", 9.2, bold=True, color=INK)
set_paragraph_format(p, after=3)

body = [
    "Respected Sir,",
    "Following our upcoming meeting, I am formally submitting this proposal for the implementation of a comprehensive Digital Institution Transformation System at Wainganga College of Engineering and Management.",
    "Operating as an autonomous multi-program institution across B.Tech, M.Tech, Polytechnic, MBA, and MCA creates a complex demand across academic, examination, finance, and compliance workflows. Our analysis indicates that the core bottleneck is not a lack of software, but the operational friction created by disconnected data systems.",
    "Our proposal is the implementation of one unified digital operating system rather than fragmented portals. The platform integrates seven institutional layers: Admissions and Conversion, Academic Operations, Examination Control, Financial Integrity, Placement and Outcomes, Governance and NAAC Compliance, and an Executive Control Dashboard.",
    "Our commitment includes high-availability cloud deployment with college-controlled data sovereignty, a perpetual exclusive institutional license model, and a discovery process focused on mapping existing institutional workflows so the software fits staff reality rather than forcing teams to adapt to generic tools.",
    "The Commercial and Operating Framework for the engagement is as follows. First, the total investment for platform design, engineering, and deployment is Rs. 6,50,000/- (Six Lakhs, Fifty Thousand Rupees only), structured as 30 percent on kickoff, 40 percent on core module delivery, and 30 percent upon final stabilization. We also provide a Lifetime Proprietary Warranty on all code against the signed blueprint, ensuring long-term institutional stability at zero additional maintenance cost.",
    "Second, the platform is engineered for cloud-native scalability, and ongoing infrastructure costs are separate from our engineering fee. These third-party operating expenses are paid directly by the college to service providers at actual usage rates and include cloud infrastructure, DLT-approved SMS and WhatsApp Business API communication layers, and annual domain and security certification costs. For an institution of this scale, we estimate these external operating costs to be approximately Rs. 18,500/- per month. SrS Logics does not mark up these costs; they are billed to the institution by the respective service providers at actuals.",
    "Upon approval, we are prepared to initiate a structured Discovery and Blueprint Phase to finalize role hierarchies, approval workflows, and the implementation sequence required for a phased go-live aligned to the upcoming academic intake.",
    "I look forward to discussing how SrS Logics can help establish a stronger, more disciplined digital foundation for Wainganga College.",
    "Sincerely,",
]

for text in body:
    add_body_paragraph(doc, text)

signature = doc.add_table(rows=1, cols=1)
signature.alignment = WD_TABLE_ALIGNMENT.LEFT
signature.autofit = False
signature.columns[0].width = Inches(3.2)
set_table_borders(signature, color=WHITE, size="0")
cell = signature.cell(0, 0)
set_cell_margins(cell, top=20, start=0, bottom=8, end=0)
set_cell_border(
    cell,
    top={"val": "single", "sz": "10", "color": NAVY},
    left={"val": "nil"},
    right={"val": "nil"},
    bottom={"val": "nil"},
)

p = cell.paragraphs[0]
r = p.add_run("Shubham Singh\n")
set_font(r, "Times New Roman", 9.8, bold=True, color=NAVY)
r = p.add_run("Founder, SrS Logics\n")
set_font(r, "Times New Roman", 8.8, bold=False, color=INK)
r = p.add_run("+91 9270925106\nshubhamsingh@srslogics.com")
set_font(r, "Times New Roman", 8.2, color=SLATE)
set_paragraph_format(p, line=0.98)

doc.save(OUT)
print(OUT)
